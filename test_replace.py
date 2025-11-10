"""
Test script to verify variable replacement in Word documents
"""
import sys
sys.path.insert(0, 'python-nlp')

from docx_service import docx_service
from io import BytesIO

# Create a simple test document
from docx import Document

doc = Document()
doc.add_paragraph("Dear [Candidate_Name],")
doc.add_paragraph("Your start date is [Start_Date] at [Company_Name].")
doc.add_paragraph("Your salary will be [Salary].")

# Save to bytes
output = BytesIO()
doc.save(output)
output.seek(0)
original_bytes = output.getvalue()

print("Original document created with variables:")
print("- [Candidate_Name]")
print("- [Start_Date]")
print("- [Company_Name]")
print("- [Salary]")
print()

# Test replacement
variables = {
    "Candidate_Name": "John Doe",
    "Start_Date": "January 15, 2025",
    "Company_Name": "Acme Corporation",
    "Salary": "$120,000"
}

print("Replacing variables:")
for key, value in variables.items():
    print(f"  {key} -> {value}")
print()

# Replace variables
modified_bytes = docx_service.replace_variables_in_docx(original_bytes, variables)

# Verify replacement
modified_doc = Document(BytesIO(modified_bytes))
full_text = '\n'.join([p.text for p in modified_doc.paragraphs])

print("Modified document content:")
print(full_text)
print()

# Check if replacements worked
all_replaced = True
for var_name in variables.keys():
    if f"[{var_name}]" in full_text:
        print(f"❌ FAILED: [{var_name}] was not replaced!")
        all_replaced = False
    else:
        print(f"✅ SUCCESS: [{var_name}] was replaced with '{variables[var_name]}'")

if all_replaced:
    print("\n🎉 All variables were successfully replaced!")
else:
    print("\n⚠️ Some variables were not replaced!")
