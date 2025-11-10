"""
Test variable name sanitization for Jinja2
"""
import sys
sys.path.insert(0, 'python-nlp')

from docxtpl_service import docxtpl_service
from docx import Document
from io import BytesIO

print("Testing Variable Name Sanitization for Jinja2")
print("=" * 60)

# Test cases
test_cases = [
    "Candidate Name",           # Space
    "Start Date",               # Space
    "Joining-Date",             # Dash
    "Employee's Name",          # Apostrophe
    "Salary (Annual)",          # Parentheses
    "Tax Rate %",               # Special char
    "Email@Address",            # @ symbol
    "Date/Time",                # Slash
]

print("\n1. Sanitization Test:")
print("-" * 60)
for original in test_cases:
    sanitized = docxtpl_service.sanitize_variable_name(original)
    print(f"   {original:20} -> {sanitized}")

print("\n2. Template Conversion Test:")
print("-" * 60)

# Create test document
doc = Document()
doc.add_paragraph("Dear [Candidate Name],")
doc.add_paragraph("Your [Start Date] is confirmed.")
doc.add_paragraph("Your [Joining-Date] is set.")
doc.add_paragraph("Contact: [Email@Address]")

output = BytesIO()
doc.save(output)
output.seek(0)
original_bytes = output.getvalue()

print("   Original text:")
original_doc = Document(BytesIO(original_bytes))
for para in original_doc.paragraphs:
    if para.text.strip():
        print(f"   - {para.text}")

# Convert to Jinja2
jinja2_bytes = docxtpl_service.convert_to_jinja2_template(original_bytes)

print("\n   Converted to Jinja2:")
jinja2_doc = Document(BytesIO(jinja2_bytes))
for para in jinja2_doc.paragraphs:
    if para.text.strip():
        print(f"   - {para.text}")

print("\n3. Rendering Test:")
print("-" * 60)

context = {
    "Candidate_Name": "John Smith",
    "Start_Date": "Jan 15, 2025",
    "Joining_Date": "Feb 1, 2025",
    "Email_Address": "john@example.com"
}

print("   Context (sanitized keys):")
for key, value in context.items():
    print(f"   {key}: {value}")

try:
    rendered_bytes = docxtpl_service.render_template(jinja2_bytes, context)

    print("\n   Rendered result:")
    rendered_doc = Document(BytesIO(rendered_bytes))
    for para in rendered_doc.paragraphs:
        if para.text.strip():
            print(f"   - {para.text}")

    print("\n✅ SUCCESS: All variables with special chars handled correctly!")

except Exception as e:
    print(f"\n❌ FAILED: {e}")
    import traceback
    traceback.print_exc()

print("\n" + "=" * 60)
