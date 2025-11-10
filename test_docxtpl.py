"""
Test docxtpl live rendering
"""
import sys
sys.path.insert(0, 'python-nlp')

from docxtpl_service import docxtpl_service
from docx import Document
from io import BytesIO

print("=" * 60)
print("Testing DocxTPL Live Rendering")
print("=" * 60)
print()

# Create a simple test document with [Variables]
doc = Document()
doc.add_paragraph("Dear [Candidate_Name],")
doc.add_paragraph("Your start date is [Start_Date] at [Company_Name].")
doc.add_paragraph("Your salary will be [Salary].")

# Save to bytes
output = BytesIO()
doc.save(output)
output.seek(0)
original_bytes = output.getvalue()

print("1. Created test document with variables:")
print("   - [Candidate_Name]")
print("   - [Start_Date]")
print("   - [Company_Name]")
print("   - [Salary]")
print()

# Step 1: Convert to Jinja2 template
print("2. Converting [Variable] to {{Variable}} (Jinja2 syntax)...")
jinja2_bytes = docxtpl_service.convert_to_jinja2_template(original_bytes)

# Check conversion
jinja2_doc = Document(BytesIO(jinja2_bytes))
jinja2_text = '\n'.join([p.text for p in jinja2_doc.paragraphs])
print("   Converted text:")
for line in jinja2_text.split('\n'):
    if line.strip():
        print(f"   {line}")
print()

# Step 2: Render with context
print("3. Rendering template with values...")
context = {
    "Candidate_Name": "Alice Johnson",
    "Start_Date": "February 1, 2025",
    "Company_Name": "Tech Innovations Inc",
    "Salary": "$135,000"
}

for key, value in context.items():
    print(f"   {key}: {value}")
print()

rendered_bytes = docxtpl_service.render_template(jinja2_bytes, context)

# Check rendered result
rendered_doc = Document(BytesIO(rendered_bytes))
rendered_text = '\n'.join([p.text for p in rendered_doc.paragraphs])

print("4. Rendered document:")
print("   " + "-" * 50)
for line in rendered_text.split('\n'):
    if line.strip():
        print(f"   {line}")
print("   " + "-" * 50)
print()

# Verify all variables were replaced
print("5. Verification:")
success = True
for var_name in context.keys():
    if f"[{var_name}]" in rendered_text or f"{{{{{var_name}}}}}" in rendered_text:
        print(f"   ❌ FAILED: {var_name} not replaced!")
        success = False
    elif context[var_name] in rendered_text:
        print(f"   ✅ SUCCESS: {var_name} -> {context[var_name]}")
    else:
        print(f"   ⚠️  WARNING: {var_name} value not found in output")
        success = False

print()
if success:
    print("🎉 DocxTPL rendering is working perfectly!")
    print("✅ Ready for live variable replacement in ONLYOFFICE")
else:
    print("⚠️  Some issues detected - check logs above")

print()
print("=" * 60)
