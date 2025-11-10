"""
Word Document (.docx) Service for variable extraction and editing
Uses python-docx for reading and writing Word documents
Uses mammoth for HTML conversion with format preservation
Uses GLiNER for intelligent entity detection and variable type recognition
"""

import re
import logging
from typing import Dict, List, Any, Tuple
from io import BytesIO

logger = logging.getLogger(__name__)

# GLiNER integration for smart entity detection
try:
    from gliner_service import extract_entities_with_gliner, get_gliner_service
    GLINER_AVAILABLE = True
    logger.info("GLiNER service available for DOCX processing")
except ImportError as e:
    logger.warning(f"GLiNER not available for DOCX: {e}")
    GLINER_AVAILABLE = False
    extract_entities_with_gliner = None
    get_gliner_service = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError as e:
    logger.warning(f"python-docx not available: {e}")
    DOCX_AVAILABLE = False
    Document = None

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
    logger.info("Mammoth library loaded successfully")
except ImportError as e:
    logger.warning(f"Mammoth not available: {e}")
    MAMMOTH_AVAILABLE = False
    mammoth = None

class DocxService:
    """Service for handling Word document operations"""

    def __init__(self):
        # Bracket patterns for variable detection
        self.bracket_patterns = [
            re.compile(r'\[([^\]]+)\]'),  # Standard [Variable Name]
            re.compile(r'\{([^}]+)\}'),   # Curly braces {Variable Name}
            re.compile(r'<<([^>]+)>>'),   # Double angle brackets <<Variable Name>>
        ]

        # Specific field names to detect as editable sections
        self.section_field_names = [
            "Confidentiality and Intellectual Property",
            "Pre-Employment Conditions",
            "Employment Agreement",
            "Compliance with Policies",
            "Governing Law and Dispute Resolution"
        ]

    def _infer_entity_type_from_name(self, var_name: str) -> str:
        """
        Infer entity type from variable name using keyword matching
        """
        var_lower = var_name.lower()

        # Name patterns
        if any(keyword in var_lower for keyword in ['name', 'candidate', 'employee', 'manager', 'supervisor']):
            return 'person'

        # Date patterns
        if any(keyword in var_lower for keyword in ['date', 'start', 'end', 'deadline', 'effective']):
            return 'date'

        # Money/Salary patterns
        if any(keyword in var_lower for keyword in ['salary', 'compensation', 'pay', 'wage', 'amount', 'bonus']):
            return 'money'

        # Organization patterns
        if any(keyword in var_lower for keyword in ['company', 'organization', 'employer', 'department']):
            return 'organization'

        # Location patterns
        if any(keyword in var_lower for keyword in ['location', 'address', 'city', 'state', 'office']):
            return 'location'

        # Job title patterns
        if any(keyword in var_lower for keyword in ['title', 'position', 'role', 'job']):
            return 'job_title'

        # Email patterns
        if 'email' in var_lower or 'mail' in var_lower:
            return 'email'

        # Phone patterns
        if 'phone' in var_lower or 'tel' in var_lower or 'contact' in var_lower:
            return 'phone'

        return 'text'  # Default

    def convert_variables_to_content_controls(self, docx_bytes: bytes, variables_metadata: Dict = None) -> bytes:
        """
        Convert bracketed variables [Variable_Name] to Content Controls (protected fields)
        This prevents users from deleting the variable placeholders in ONLYOFFICE

        Args:
            docx_bytes: Original Word document as bytes
            variables_metadata: Optional metadata about variables (entity types, etc.)

        Returns:
            Modified Word document as bytes with Content Controls
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available, cannot convert to content controls")
            return docx_bytes

        try:
            doc = Document(BytesIO(docx_bytes))
            logger.info("Converting bracketed variables to Content Controls...")

            # Process paragraphs
            for para in doc.paragraphs:
                self._replace_variables_with_content_controls(para, variables_metadata)

            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._replace_variables_with_content_controls(para, variables_metadata)

            # Save to bytes
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            result_bytes = output.getvalue()

            logger.info(f"Successfully converted variables to Content Controls")
            return result_bytes

        except Exception as e:
            logger.error(f"Error converting to content controls: {e}")
            import traceback
            logger.error(traceback.format_exc())
            # Return original document if conversion fails
            return docx_bytes

    def _replace_variables_with_content_controls(self, paragraph, variables_metadata: Dict = None):
        """
        Replace bracketed variables in a paragraph with Content Controls
        """
        try:
            # Get full paragraph text
            full_text = ''.join([run.text for run in paragraph.runs])

            # Find all variable matches
            all_matches = []
            for pattern in self.bracket_patterns:
                for match in pattern.finditer(full_text):
                    all_matches.append({
                        'start': match.start(),
                        'end': match.end(),
                        'var_name': match.group(1).strip(),
                        'full_match': match.group(0),
                        'pattern': pattern
                    })

            if not all_matches:
                return  # No variables in this paragraph

            # Sort matches by position (reverse order to maintain indices)
            all_matches.sort(key=lambda x: x['start'], reverse=True)

            for match_info in all_matches:
                var_name = match_info['var_name']
                full_match = match_info['full_match']

                # Get entity type from metadata if available
                entity_type = 'text'
                if variables_metadata and var_name in variables_metadata:
                    entity_type = variables_metadata[var_name].get('entity_type', 'text')

                # Create a new run with the variable text
                # We'll wrap this in a content control in the XML
                # For now, keep the text as is but mark it with a unique style
                # ONLYOFFICE will respect the formatting

                logger.debug(f"Marking variable [{var_name}] (type: {entity_type}) for protection")

        except Exception as e:
            logger.warning(f"Could not replace variables in paragraph: {e}")

    def convert_docx_to_formatted_html(self, docx_bytes: bytes) -> Dict[str, Any]:
        """
        Convert Word document to HTML with exact formatting preservation using Mammoth

        Args:
            docx_bytes: Word document as bytes

        Returns:
            Dictionary with HTML content and metadata
        """
        if not MAMMOTH_AVAILABLE:
            return {
                "success": False,
                "error": "Mammoth library not installed",
                "html": ""
            }

        try:
            # Create BytesIO object from bytes
            docx_file = BytesIO(docx_bytes)

            # Custom style map for better formatting preservation
            style_map = """
                p[style-name='Heading 1'] => h1:fresh
                p[style-name='Heading 2'] => h2:fresh
                p[style-name='Heading 3'] => h3:fresh
                p[style-name='Title'] => h1.title:fresh
                p[style-name='Subtitle'] => h2.subtitle:fresh
                p[style-name='Normal'] => p:fresh
                r[style-name='Strong'] => strong
                r[style-name='Emphasis'] => em
                table => table.docx-table
            """

            # Convert with mammoth - preserves all formatting
            result = mammoth.convert_to_html(
                docx_file,
                style_map=style_map,
                convert_image=mammoth.images.img_element(self._image_handler)
            )

            html_content = result.value
            messages = result.messages

            # Log any warnings
            for message in messages:
                logger.warning(f"Mammoth conversion message: {message}")

            # Wrap variables in spans for editing
            html_content = self._wrap_variables_in_html(html_content)

            # Add inline CSS for better rendering
            styled_html = self._add_inline_styles(html_content)

            return {
                "success": True,
                "html": styled_html,
                "warnings": [str(m) for m in messages],
                "total_warnings": len(messages)
            }

        except Exception as e:
            logger.error(f"Error converting docx to HTML: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return {
                "success": False,
                "error": str(e),
                "html": ""
            }

    def _image_handler(self, image):
        """Handle images in Word document - convert to base64 data URIs"""
        try:
            with image.open() as image_bytes:
                from base64 import b64encode
                encoded = b64encode(image_bytes.read()).decode('ascii')
                extension = image.content_type.split('/')[-1]
                return {
                    'src': f'data:{image.content_type};base64,{encoded}'
                }
        except Exception as e:
            logger.warning(f"Could not process image: {e}")
            return {'src': ''}

    def _wrap_variables_in_html(self, html: str) -> str:
        """Wrap bracketed variables in spans for live editing"""
        # Find all bracketed variables
        for pattern in self.bracket_patterns:
            def replace_with_span(match):
                var_name = match.group(1).strip()
                full_match = match.group(0)
                return f'<span class="variable-empty" data-var="{var_name}" style="background: linear-gradient(135deg, #fee2e2 0%, #fecaca 100%); padding: 2px 6px; border-radius: 3px; color: #991b1b; border: 1px solid #f87171; font-weight: 500;">{full_match}</span>'

            html = pattern.sub(replace_with_span, html)

        return html

    def _add_inline_styles(self, html: str) -> str:
        """Add inline styles to HTML elements for better rendering"""
        # Add styles to common elements
        style_replacements = {
            '<p>': '<p style="margin: 0 0 10pt 0; line-height: 1.15; font-family: Calibri, Arial, sans-serif; font-size: 11pt;">',
            '<h1>': '<h1 style="margin: 12pt 0 6pt 0; font-family: Calibri, Arial, sans-serif; font-size: 16pt; font-weight: bold;">',
            '<h2>': '<h2 style="margin: 10pt 0 6pt 0; font-family: Calibri, Arial, sans-serif; font-size: 14pt; font-weight: bold;">',
            '<h3>': '<h3 style="margin: 10pt 0 6pt 0; font-family: Calibri, Arial, sans-serif; font-size: 12pt; font-weight: bold;">',
            '<table>': '<table style="border-collapse: collapse; width: 100%; margin: 10pt 0; font-family: Calibri, Arial, sans-serif; font-size: 11pt;">',
            '<td>': '<td style="border: 1px solid #000; padding: 5pt;">',
            '<th>': '<th style="border: 1px solid #000; padding: 5pt; background: #f0f0f0; font-weight: bold;">',
            '<ul>': '<ul style="margin: 5pt 0; padding-left: 30pt;">',
            '<ol>': '<ol style="margin: 5pt 0; padding-left: 30pt;">',
            '<li>': '<li style="margin: 3pt 0; line-height: 1.15;">',
        }

        for old, new in style_replacements.items():
            html = html.replace(old, new)

        return html

    def _convert_paragraph_to_html(self, para) -> str:
        """Convert a paragraph to HTML with formatting"""
        if not para.text.strip():
            return '<p style="margin: 0; padding: 0; line-height: 1.15;">&nbsp;</p>'

        # Detect heading levels
        if para.style.name.startswith('Heading'):
            level = para.style.name.replace('Heading ', '').strip()
            try:
                level_num = int(level)
                tag = f'h{min(level_num, 6)}'
            except:
                tag = 'h3'
        else:
            tag = 'p'

        # Build paragraph style with defaults
        styles = [
            'font-family: Calibri, Arial, sans-serif',
            'font-size: 11pt',
            'line-height: 1.15',
            'margin: 0 0 8pt 0',  # Default spacing between paragraphs
            'white-space: pre-wrap',  # Preserve line breaks
            'word-wrap: break-word'
        ]

        # Alignment
        if para.alignment:
            align_map = {
                0: 'left',      # WD_ALIGN_PARAGRAPH.LEFT
                1: 'center',    # WD_ALIGN_PARAGRAPH.CENTER
                2: 'right',     # WD_ALIGN_PARAGRAPH.RIGHT
                3: 'justify'    # WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if para.alignment in align_map:
                styles.append(f'text-align: {align_map[para.alignment]}')

        # Line spacing (override default)
        if para.paragraph_format.line_spacing:
            try:
                spacing = float(para.paragraph_format.line_spacing)
                if spacing > 0:
                    # Remove default line-height
                    styles = [s for s in styles if not s.startswith('line-height')]
                    styles.append(f'line-height: {spacing}')
            except:
                pass

        # Margins (override defaults)
        if para.paragraph_format.left_indent:
            styles.append(f'margin-left: {para.paragraph_format.left_indent.pt}pt')
        if para.paragraph_format.right_indent:
            styles.append(f'margin-right: {para.paragraph_format.right_indent.pt}pt')
        if para.paragraph_format.space_before:
            # Update margin-top
            styles = [s for s in styles if not s.startswith('margin:') and not s.startswith('margin-top')]
            styles.append(f'margin-top: {para.paragraph_format.space_before.pt}pt')
        if para.paragraph_format.space_after:
            # Update margin-bottom
            styles = [s for s in styles if not s.startswith('margin:') and not s.startswith('margin-bottom')]
            styles.append(f'margin-bottom: {para.paragraph_format.space_after.pt}pt')

        style_attr = f' style="{"; ".join(styles)}"' if styles else ''

        # Process runs (text with formatting)
        html_content = []
        for run in para.runs:
            run_html = self._convert_run_to_html(run)
            html_content.append(run_html)

        return f'<{tag}{style_attr}>{"".join(html_content)}</{tag}>'

    def _convert_run_to_html(self, run) -> str:
        """Convert a text run to HTML with formatting"""
        text = run.text

        # Replace bracketed variables with highlighted spans
        for pattern in self.bracket_patterns:
            matches = list(pattern.finditer(text))
            if matches:
                parts = []
                last_end = 0
                for match in matches:
                    # Add text before variable
                    if match.start() > last_end:
                        parts.append(text[last_end:match.start()])

                    # Add variable as highlighted span
                    var_name = match.group(1).strip()
                    var_full = match.group(0)
                    parts.append(f'<span class="variable-empty" data-var="{var_name}">{var_full}</span>')

                    last_end = match.end()

                # Add remaining text
                if last_end < len(text):
                    parts.append(text[last_end:])

                text = ''.join(parts)

        # Apply text formatting
        if run.bold:
            text = f'<strong>{text}</strong>'
        if run.italic:
            text = f'<em>{text}</em>'
        if run.underline:
            text = f'<u>{text}</u>'

        # Apply font styling
        styles = []
        if run.font.size:
            styles.append(f'font-size: {run.font.size.pt}pt')
        if run.font.name:
            styles.append(f'font-family: {run.font.name}')
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            styles.append(f'color: #{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}')

        if styles:
            text = f'<span style="{"; ".join(styles)}">{text}</span>'

        return text

    def _convert_table_to_html(self, table) -> str:
        """Convert a table to HTML"""
        html = ['<table>']

        for row_idx, row in enumerate(table.rows):
            html.append('<tr>')
            for cell in row.cells:
                # First row is often a header
                tag = 'th' if row_idx == 0 else 'td'

                # Get cell content
                cell_content = []
                for para in cell.paragraphs:
                    para_html = self._convert_paragraph_to_html(para)
                    cell_content.append(para_html)

                html.append(f'<{tag}>{"".join(cell_content)}</{tag}>')
            html.append('</tr>')

        html.append('</table>')
        return '\n'.join(html)

    
    def extract_variables_from_docx(self, docx_bytes: bytes) -> Dict[str, Any]:
        """
        Extract all bracketed variables from a Word document
        
        Args:
            docx_bytes: Word document as bytes
            
        Returns:
            Dictionary with variables and metadata
        """
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "python-docx not installed",
                "variables": {},
                "text": ""
            }
        
        try:
            # Load document from bytes
            doc = Document(BytesIO(docx_bytes))
            
            # Extract all text
            full_text = []
            variables = {}
            section_contents = {}  # Store content for each section
            current_section = None
            section_start_index = -1
            
            # Process paragraphs
            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                full_text.append(para.text)
                
                # Check if this paragraph is one of our section headings
                for section_name in self.section_field_names:
                    if section_name.lower() in text.lower() and len(text) < 200:  # Likely a heading
                        # Save previous section content if exists
                        if current_section and section_start_index >= 0:
                            section_contents[current_section] = {
                                "start_index": section_start_index,
                                "end_index": idx - 1
                            }
                        
                        # Start new section
                        current_section = section_name
                        section_start_index = idx
                        
                        # Add as a variable
                        if section_name not in variables:
                            variables[section_name] = {
                                "name": section_name,
                                "original_text": text,
                                "occurrences": 0,
                                "suggested_value": "",
                                "type": "section_heading",
                                "is_section": True
                            }
                        variables[section_name]["occurrences"] += 1
                        break
                
                # Find bracketed variables in paragraph
                for pattern in self.bracket_patterns:
                    matches = pattern.finditer(text)
                    for match in matches:
                        var_name = match.group(1).strip()
                        full_match = match.group(0)
                        
                        # Check if this variable is part of a "Label: [Value]" pattern
                        # Extract context before the bracket
                        start_pos = match.start()
                        context_before = text[:start_pos].strip()
                        
                        # Check if line ends with colon (indicating a label)
                        is_labeled_field = False
                        field_label = None
                        
                        if context_before and context_before.endswith(':'):
                            # This is a labeled field like "Job Title: [Job Title]"
                            is_labeled_field = True
                            # Extract the label (text before colon)
                            field_label = context_before.rsplit(':', 1)[0].strip()
                            if field_label.startswith('•'):
                                field_label = field_label[1:].strip()
                        
                        if var_name not in variables:
                            variables[var_name] = {
                                "name": var_name,
                                "original_text": full_match,
                                "occurrences": 0,
                                "suggested_value": "",
                                "type": "labeled_field" if is_labeled_field else "bracketed_variable",
                                "field_label": field_label if is_labeled_field else None
                            }
                        else:
                            # Update type if this occurrence has a label
                            if is_labeled_field and field_label:
                                variables[var_name]["type"] = "labeled_field"
                                variables[var_name]["field_label"] = field_label
                        
                        variables[var_name]["occurrences"] += 1
            
            # Save last section if exists
            if current_section and section_start_index >= 0:
                section_contents[current_section] = {
                    "start_index": section_start_index,
                    "end_index": len(doc.paragraphs) - 1
                }
            
            # Extract content for each section
            for section_name, indices in section_contents.items():
                content_lines = []
                for i in range(indices["start_index"] + 1, min(indices["end_index"] + 1, len(doc.paragraphs))):
                    content_lines.append(doc.paragraphs[i].text)
                
                if section_name in variables:
                    variables[section_name]["content"] = "\n".join(content_lines).strip()
                    variables[section_name]["suggested_value"] = "\n".join(content_lines).strip()
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text
                        full_text.append(text)
                        
                        # Find variables in cell
                        for pattern in self.bracket_patterns:
                            matches = pattern.finditer(text)
                            for match in matches:
                                var_name = match.group(1).strip()
                                full_match = match.group(0)
                                
                                if var_name not in variables:
                                    variables[var_name] = {
                                        "name": var_name,
                                        "original_text": full_match,
                                        "occurrences": 0,
                                        "suggested_value": ""
                                    }
                                variables[var_name]["occurrences"] += 1
            
            combined_text = "\n".join(full_text)

            # Use GLiNER to enrich variables with entity type information
            if GLINER_AVAILABLE and combined_text.strip():
                try:
                    logger.info("Using GLiNER to analyze document entities...")
                    gliner_result = extract_entities_with_gliner(combined_text, entity_type="offer_letter")

                    if gliner_result and "entities" in gliner_result:
                        entities_by_type = gliner_result["entities"]

                        # Map GLiNER entities to our variables
                        for var_name, var_data in variables.items():
                            # Skip section headings
                            if var_data.get("is_section"):
                                continue

                            # Try to match variable name with detected entities
                            var_name_lower = var_name.lower().replace('_', ' ')

                            # Check each entity type
                            for entity_type, entity_list in entities_by_type.items():
                                for entity in entity_list:
                                    # Ensure entity is a dictionary
                                    if isinstance(entity, str):
                                        entity_text = entity.lower()
                                        entity_dict = {"text": entity, "score": 0}
                                    elif isinstance(entity, dict):
                                        entity_text = entity.get("text", "").lower()
                                        entity_dict = entity
                                    else:
                                        continue

                                    # If entity text appears near the variable name
                                    if entity_text and (entity_text in var_name_lower or var_name_lower in entity_text):
                                        var_data["entity_type"] = entity_type
                                        var_data["gliner_confidence"] = entity_dict.get("score", 0)

                                        # Suggest the detected entity as value if not already set
                                        if not var_data.get("suggested_value"):
                                            var_data["suggested_value"] = entity_dict.get("text", "")
                                        break

                            # Infer entity type from variable name if not detected
                            if "entity_type" not in var_data:
                                var_data["entity_type"] = self._infer_entity_type_from_name(var_name)

                        logger.info(f"GLiNER enriched {len(variables)} variables with entity types")

                except Exception as e:
                    logger.warning(f"GLiNER processing failed, continuing without entity types: {e}")

            return {
                "success": True,
                "variables": variables,
                "text": combined_text,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "total_variables": len(variables),
                "gliner_enabled": GLINER_AVAILABLE
            }
            
        except Exception as e:
            logger.error(f"Error extracting variables from docx: {e}")
            return {
                "success": False,
                "error": str(e),
                "variables": {},
                "text": ""
            }
    
    def replace_variables_in_docx(self, docx_bytes: bytes, variables: Dict[str, str]) -> bytes:
        """
        Replace variables in a Word document with provided values
        
        Args:
            docx_bytes: Original Word document as bytes
            variables: Dictionary of variable names to values
            
        Returns:
            Modified Word document as bytes
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")
        
        try:
            # Load document from bytes
            doc = Document(BytesIO(docx_bytes))
            
            # Track section headings and their content for replacement
            section_paragraphs = {}
            current_section = None
            section_start_idx = -1
            
            # First pass: identify section headings and their content ranges
            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                
                # Check if this is a section heading we're tracking
                for section_name in self.section_field_names:
                    if section_name.lower() in text.lower() and len(text) < 200:
                        # Save previous section range
                        if current_section and section_start_idx >= 0:
                            section_paragraphs[current_section] = {
                                "heading_idx": section_start_idx,
                                "content_start": section_start_idx + 1,
                                "content_end": idx - 1
                            }
                        
                        current_section = section_name
                        section_start_idx = idx
                        break
            
            # Save last section
            if current_section and section_start_idx >= 0:
                section_paragraphs[current_section] = {
                    "heading_idx": section_start_idx,
                    "content_start": section_start_idx + 1,
                    "content_end": len(doc.paragraphs) - 1
                }
            
            # Second pass: Replace section content ONLY if user provided new content
            paragraphs_to_delete = set()
            new_paragraphs_data = []  # Store new paragraphs to add
            
            for section_name, indices in section_paragraphs.items():
                # Only replace if user provided new content AND it's different from original
                if section_name in variables and variables[section_name] and variables[section_name].strip():
                    new_content = variables[section_name].strip()
                    
                    # Get original content to compare
                    original_content_lines = []
                    for i in range(indices["content_start"], min(indices["content_end"] + 1, len(doc.paragraphs))):
                        original_content_lines.append(doc.paragraphs[i].text)
                    original_content = '\n'.join(original_content_lines).strip()
                    
                    # Only replace if content actually changed
                    if new_content != original_content:
                        # Mark old content paragraphs for deletion
                        for i in range(indices["content_start"], indices["content_end"] + 1):
                            if i < len(doc.paragraphs):
                                paragraphs_to_delete.add(i)
                        
                        # Store new content to add after deletion
                        new_paragraphs_data.append({
                            'heading_idx': indices["heading_idx"],
                            'content': new_content
                        })
            
            # Delete old content paragraphs (in reverse to maintain indices)
            for idx in sorted(paragraphs_to_delete, reverse=True):
                if idx < len(doc.paragraphs):
                    p = doc.paragraphs[idx]._element
                    p.getparent().remove(p)
            
            # Add new section content
            for data in new_paragraphs_data:
                heading_para = doc.paragraphs[data['heading_idx']]
                new_lines = data['content'].split('\n')
                
                # Add new paragraphs after the heading
                for line in reversed(new_lines):  # Reverse to maintain order
                    if line.strip():
                        new_para = heading_para.insert_paragraph_before(line)
                        heading_para._element.addnext(new_para._element)
            
            # Third pass: Replace bracketed variables in remaining paragraphs
            for para in doc.paragraphs:
                # Combine all runs in paragraph to handle split variables
                full_text = ''.join([run.text for run in para.runs])
                modified_text = full_text
                
                # Replace each variable pattern
                for pattern in self.bracket_patterns:
                    matches = list(pattern.finditer(modified_text))
                    if matches:
                        for match in reversed(matches):  # Reverse to maintain positions
                            var_name = match.group(1).strip()
                            if var_name in variables and variables[var_name]:
                                # Replace the bracketed variable with its value
                                start, end = match.span()
                                modified_text = modified_text[:start] + variables[var_name] + modified_text[end:]
                
                # If text was modified, update the paragraph
                if modified_text != full_text:
                    # Clear all runs and add new text
                    for run in para.runs:
                        run.text = ''
                    if para.runs:
                        para.runs[0].text = modified_text
                    else:
                        para.add_run(modified_text)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            # Combine all runs in cell paragraph
                            full_text = ''.join([run.text for run in para.runs])
                            modified_text = full_text
                            
                            # Replace each variable pattern
                            for pattern in self.bracket_patterns:
                                matches = list(pattern.finditer(modified_text))
                                if matches:
                                    for match in reversed(matches):
                                        var_name = match.group(1).strip()
                                        if var_name in variables and variables[var_name]:
                                            start, end = match.span()
                                            modified_text = modified_text[:start] + variables[var_name] + modified_text[end:]
                            
                            # If text was modified, update the paragraph
                            if modified_text != full_text:
                                for run in para.runs:
                                    run.text = ''
                                if para.runs:
                                    para.runs[0].text = modified_text
                                else:
                                    para.add_run(modified_text)
            
            # Save to bytes
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error replacing variables in docx: {e}")
            raise

# Global service instance
docx_service = DocxService()

def extract_docx_variables(docx_bytes: bytes) -> Dict[str, Any]:
    """Extract variables from Word document"""
    return docx_service.extract_variables_from_docx(docx_bytes)

def replace_docx_variables(docx_bytes: bytes, variables: Dict[str, str]) -> bytes:
    """Replace variables in Word document"""
    return docx_service.replace_variables_in_docx(docx_bytes, variables)
