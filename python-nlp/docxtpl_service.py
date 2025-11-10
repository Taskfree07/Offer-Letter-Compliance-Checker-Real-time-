"""
Document Template Service using docxtpl
This service provides LIVE variable replacement in Word documents
"""
import re
import logging
from typing import Dict, Any
from io import BytesIO
from docxtpl import DocxTemplate

logger = logging.getLogger(__name__)

class DocxTemplateService:
    """Service for handling Word document templates with live rendering"""

    def __init__(self):
        # Patterns to detect different variable formats
        self.bracket_patterns = [
            re.compile(r'\[([^\]]+)\]'),      # [Variable]
            re.compile(r'\{\{([^}]+)\}\}'),   # {{Variable}} - Jinja2 style
            re.compile(r'\{([^}]+)\}'),       # {Variable}
        ]

    def sanitize_variable_name(self, var_name: str) -> str:
        """
        Sanitize variable name for Jinja2 compatibility
        Replace spaces and special chars with underscores
        """
        # Remove leading/trailing whitespace
        var_name = var_name.strip()

        # Replace spaces with underscores
        var_name = re.sub(r'\s+', '_', var_name)

        # Remove any characters that aren't alphanumeric or underscore
        var_name = re.sub(r'[^\w]', '_', var_name)

        # Remove multiple consecutive underscores
        var_name = re.sub(r'_+', '_', var_name)

        # Remove leading/trailing underscores
        var_name = var_name.strip('_')

        return var_name

    def convert_to_jinja2_template(self, docx_bytes: bytes) -> bytes:
        """
        Convert [Variable] placeholders to {{Variable}} Jinja2 syntax
        This allows docxtpl to render them dynamically
        """
        try:
            from docx import Document

            doc = Document(BytesIO(docx_bytes))
            logger.info("Converting bracketed variables to Jinja2 syntax...")

            def convert_text(text):
                """Convert [Variable Name] to {{ Variable_Name }}"""
                def replace_var(match):
                    var_name = match.group(1).strip()
                    sanitized = self.sanitize_variable_name(var_name)
                    return f'{{{{ {sanitized} }}}}'

                return re.sub(r'\[([^\]]+)\]', replace_var, text)

            # Convert in paragraphs
            for para in doc.paragraphs:
                if '[' in para.text:
                    # Get all runs text
                    full_text = ''.join([run.text for run in para.runs])

                    # Convert [Variable] to {{ Variable }}
                    converted_text = convert_text(full_text)

                    if converted_text != full_text:
                        # Clear all runs and set new text
                        for run in para.runs:
                            run.text = ''
                        if para.runs:
                            para.runs[0].text = converted_text
                        logger.debug(f"Converted: {full_text[:50]}... -> {converted_text[:50]}...")

            # Convert in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if '[' in para.text:
                                full_text = ''.join([run.text for run in para.runs])
                                converted_text = convert_text(full_text)

                                if converted_text != full_text:
                                    for run in para.runs:
                                        run.text = ''
                                    if para.runs:
                                        para.runs[0].text = converted_text

            # Save to bytes
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            result_bytes = output.getvalue()

            logger.info(f"✅ Converted document to Jinja2 template format")
            return result_bytes

        except Exception as e:
            logger.error(f"Error converting to Jinja2: {e}")
            import traceback
            logger.error(traceback.format_exc())
            # Return original if conversion fails
            return docx_bytes

    def extract_variables(self, docx_bytes: bytes) -> Dict[str, Any]:
        """
        Extract variables from document (supports both [Variable] and {{Variable}})
        """
        try:
            from docx import Document

            doc = Document(BytesIO(docx_bytes))
            variables = {}

            # Extract from paragraphs
            for para in doc.paragraphs:
                text = para.text

                # Find [Variable] format
                for match in re.finditer(r'\[([^\]]+)\]', text):
                    var_name = match.group(1).strip()
                    if var_name not in variables:
                        variables[var_name] = {
                            "name": var_name,
                            "value": "",
                            "occurrences": 0
                        }
                    variables[var_name]["occurrences"] += 1

                # Find {{Variable}} format
                for match in re.finditer(r'\{\{([^}]+)\}\}', text):
                    var_name = match.group(1).strip()
                    if var_name not in variables:
                        variables[var_name] = {
                            "name": var_name,
                            "value": "",
                            "occurrences": 0
                        }
                    variables[var_name]["occurrences"] += 1

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text

                        for match in re.finditer(r'\[([^\]]+)\]|\{\{([^}]+)\}\}', text):
                            var_name = (match.group(1) or match.group(2)).strip()
                            if var_name not in variables:
                                variables[var_name] = {
                                    "name": var_name,
                                    "value": "",
                                    "occurrences": 0
                                }
                            variables[var_name]["occurrences"] += 1

            logger.info(f"📋 Extracted {len(variables)} variables")
            return {
                "success": True,
                "variables": variables,
                "total_variables": len(variables)
            }

        except Exception as e:
            logger.error(f"Error extracting variables: {e}")
            return {
                "success": False,
                "error": str(e),
                "variables": {}
            }

    def render_template(self, docx_bytes: bytes, context: Dict[str, str]) -> bytes:
        """
        Render template with given context using docxtpl
        This is the LIVE rendering that shows changes immediately
        """
        try:
            # Create DocxTemplate from bytes
            template = DocxTemplate(BytesIO(docx_bytes))

            logger.info(f"🎨 Rendering template with {len(context)} variables")

            # Render with context
            template.render(context)

            # Save to bytes
            output = BytesIO()
            template.save(output)
            output.seek(0)
            result_bytes = output.getvalue()

            logger.info(f"✅ Template rendered successfully ({len(result_bytes)} bytes)")
            return result_bytes

        except Exception as e:
            logger.error(f"Error rendering template: {e}")
            import traceback
            logger.error(traceback.format_exc())
            raise

    def render_template_live(self, docx_bytes: bytes, context: Dict[str, str]) -> bytes:
        """
        Live rendering - converts and renders in one step
        This is called every time a variable changes
        """
        try:
            # First, ensure document uses Jinja2 syntax
            jinja2_bytes = self.convert_to_jinja2_template(docx_bytes)

            # Then render with context
            rendered_bytes = self.render_template(jinja2_bytes, context)

            return rendered_bytes

        except Exception as e:
            logger.error(f"Error in live rendering: {e}")
            raise


# Global service instance
docxtpl_service = DocxTemplateService()


def extract_template_variables(docx_bytes: bytes) -> Dict[str, Any]:
    """Extract variables from template"""
    return docxtpl_service.extract_variables(docx_bytes)


def render_docx_template(docx_bytes: bytes, context: Dict[str, str]) -> bytes:
    """Render template with context"""
    return docxtpl_service.render_template(docx_bytes, context)


def render_docx_template_live(docx_bytes: bytes, context: Dict[str, str]) -> bytes:
    """Live render - convert and render in one step"""
    return docxtpl_service.render_template_live(docx_bytes, context)
